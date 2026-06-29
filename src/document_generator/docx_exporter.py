"""
Экспорт заполненного документа в формат DOCX с правильным форматированием.
"""
import os
import re
from datetime import datetime
from typing import Dict, Optional
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


TEMPLATES_DIR = "./templates"
OUTPUT_DIR = "./output/documents"


def ensure_output_dir():
    """Создаёт директорию для выходных файлов, если её нет."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)


def load_template_text(doc_type: str) -> Optional[str]:
    """Загружает текстовый шаблон."""
    path = os.path.join(TEMPLATES_DIR, f"{doc_type}.txt")
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def fill_template(template: str, fields: Dict[str, str]) -> str:
    """Заполняет шаблон значениями полей."""
    result = template
    for key, value in fields.items():
        if value:
            result = result.replace(f"[{key}]", str(value))
    
    # Заменяем оставшиеся незаполненные плейсхолдеры на прочерки
    result = re.sub(r'\[([^\]]+)\]', '___________', result)
    
    # Если в шаблоне была дата и она не заполнена — ставим текущую
    if "___________" in result and "[date]" in template:
        today = datetime.now().strftime("%d.%m.%Y")
        # Заменяем только первое вхождение прочерков (которое было датой)
        result = result.replace("___________", today, 1)
    
    return result


def create_docx(text: str, doc_type: str, fields: Dict[str, str]) -> str:
    """
    Создаёт DOCX документ из заполненного текста.
    
    Args:
        text: Заполненный текст документа
        doc_type: Тип документа
        fields: Словарь с полями
    
    Returns:
        Путь к созданному DOCX файлу
    """
    ensure_output_dir()
    
    doc = Document()
    
    # Настройка полей страницы (стандарт для юридических документов)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    
    # Устанавливаем шрифт по умолчанию (Times New Roman 14)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    # Разбиваем текст на абзацы
    paragraphs = text.split('\n')
    
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        
        paragraph = doc.add_paragraph()
        
        # Определяем тип абзаца и форматируем
        if _is_header(para_text, fields):
            # Шапка документа — по правому краю
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run(para_text)
            run.font.size = Pt(14)
        elif _is_title(para_text):
            # Заголовок — по центру, жирный
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(para_text)
            run.bold = True
            run.font.size = Pt(14)
        elif _is_signature_line(para_text):
            # Строка подписи
            run = paragraph.add_run(para_text)
            run.font.size = Pt(14)
        elif _is_list_item(para_text):
            # Пункты списка
            run = paragraph.add_run(para_text)
            run.font.size = Pt(14)
        else:
            # Обычный абзац — по ширине с отступом
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            run = paragraph.add_run(para_text)
            run.font.size = Pt(14)
    
    # Формируем имя файла
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{doc_type}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    
    doc.save(filepath)
    
    print(f"✅ DOCX создан: {filepath}")
    return filepath


def _is_header(text: str, fields: Dict[str, str]) -> bool:
    """Определяет, является ли абзац шапкой документа."""
    header_keywords = ['В ', 'в ', 'Истец:', 'Ответчик:', 'Заявитель:', 
                      'адрес:', 'по делу №', 'судья']
    
    if any(text.startswith(kw) for kw in header_keywords):
        return True
    
    court_name = fields.get('court_name') or fields.get('appeal_court_name')
    plaintiff = fields.get('plaintiff_name') or fields.get('applicant_name')
    defendant = fields.get('defendant_name')
    
    if court_name and court_name in text:
        return True
    if plaintiff and plaintiff in text and len(text) < 150:
        return True
    if defendant and defendant in text and len(text) < 150:
        return True
    
    return False


def _is_title(text: str) -> bool:
    """Определяет, является ли абзац заголовком."""
    title_keywords = [
        'Исковое заявление', 'Заявление', 'Жалоба', 'Возражение',
        'Ходатайство', 'Апелляционная', 'Кассационная', 'Претензия'
    ]
    return any(kw in text for kw in title_keywords) and len(text) < 150


def _is_signature_line(text: str) -> bool:
    """Определяет, является ли абзац строкой подписи."""
    if re.match(r'^\d{2}\.\d{2}\.\d{4}', text):
        return True
    if '_________________' in text or '_______________' in text:
        return True
    return False


def _is_list_item(text: str) -> bool:
    """Определяет, является ли абзац пунктом списка."""
    if re.match(r'^\d+[\.\)]\s', text):
        return True
    if re.match(r'^[а-я]\)\s', text):
        return True
    if re.match(r'^[IVX]+\.\s', text):
        return True
    return False


def test_docx_generation():
    """Тест генерации DOCX."""
    print("🧪 ТЕСТ ГЕНЕРАЦИИ DOCX")
    print("="*60)
    
    # Тестовые поля
    fields = {
        "court_name": "Ленинский районный суд г. Кызыла",
        "plaintiff_name": "Петров Петр Петрович",
        "plaintiff_address": "г. Кызыл, ул. Титова, д. 10, кв. 5",
        "defendant_name": "ООО «Технопарк»",
        "defendant_address": "г. Кызыл, ул. Ленина, д. 10",
        "product_name": "Samsung Galaxy S23",
        "amount": "80000",
        "defect_description": "мерцание экрана через неделю после покупки",
        "moral_damage": "10000",
        "date": "29.06.2026"
    }
    
    # Загружаем шаблон
    template = load_template_text("claim_consumer")
    if not template:
        print("❌ Шаблон claim_consumer не найден")
        return
    
    # Заполняем шаблон
    filled_text = fill_template(template, fields)
    print(f"\n📄 Заполненный текст:\n{filled_text}\n")
    
    # Создаём DOCX
    filepath = create_docx(filled_text, "claim_consumer", fields)
    
    print(f"\n✅ Файл создан: {filepath}")
    print(f"   Размер: {os.path.getsize(filepath)} байт")


if __name__ == "__main__":
    test_docx_generation()
